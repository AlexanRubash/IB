from docx import Document
from docx.shared import RGBColor

def text_to_binary(text):
    """Преобразует текст в двоичный формат"""
    return ''.join(format(ord(char), '08b') for char in text)

def binary_to_text(binary):
    """Преобразует двоичный текст обратно в текст"""
    chars = [binary[i:i+8] for i in range(0, len(binary), 8)]
    return ''.join(chr(int(char, 2)) for char in chars)

def modify_color(color, bit):
    """Изменяет младший бит синего компонента цвета"""
    # color - это (R, G, B)
    # bit - это 0 или 1
    new_blue = (color[2] & 0xFE) | bit  # Изменяем младший бит синего компонента
    return RGBColor(color[0], color[1], new_blue)

def encode_message_in_text_color(doc_path, message, output_path):
    """Кодирует сообщение в цвет отдельных символов текста"""
    doc = Document(doc_path)
    binary_message = text_to_binary(message)
    message_length = len(binary_message)
    message_index = 0

    for para in doc.paragraphs:
        for run in para.runs:
            new_text = ''
            for char in run.text:
                if message_index < message_length:
                    color = run.font.color.rgb if run.font.color else RGBColor(0, 0, 0)
                    new_color = modify_color(color, int(binary_message[message_index]))
                    new_run = para.add_run(char)
                    new_run.font.color.rgb = new_color
                    # Сохраняем остальные атрибуты шрифта
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    new_run.font.underline = run.font.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    message_index += 1
                else:
                    new_run = para.add_run(char)
                    new_run.font.color.rgb = run.font.color.rgb if run.font.color else RGBColor(0, 0, 0)
                    # Сохраняем остальные атрибуты шрифта
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
                    new_run.font.underline = run.font.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name

            run.clear()  # Очищаем старый run

    doc.save(output_path)

def decode_message_from_text_color(doc_path, message_length):
    """Декодирует сообщение из цвета отдельных символов текста"""
    doc = Document(doc_path)
    binary_message = []

    for para in doc.paragraphs:
        for run in para.runs:
            for char in run.text:
                if len(binary_message) < message_length * 8:
                    if run.font.color:
                        color = run.font.color.rgb
                        binary_message.append(str(color[2] & 1))

    binary_message = ''.join(binary_message)
    return binary_to_text(binary_message[:message_length * 8])

# Пример использования
input_doc_path = 'input.docx'
output_doc_path = 'output.docx'
message = 'Hello, World!'

encode_message_in_text_color(input_doc_path, message, output_doc_path)
decoded_message = decode_message_from_text_color(output_doc_path, len(message))
print('Decoded message:', decoded_message)
